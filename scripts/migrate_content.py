"""Create the static site content directly from the original Material files.

The generated MDX records retain the original subject, week, title, text and
source path. No old generated answer or code artifact is used as a source.
"""

from __future__ import annotations

import json
import re
import shutil
import unicodedata
from collections import defaultdict
from pathlib import Path

from docx import Document


PROJECT = Path(__file__).resolve().parents[1]
SOURCE = PROJECT.parent / "Material"
CONTENT = PROJECT / "src" / "content" / "aulas"
PUBLIC_ARTIFACTS = PROJECT / "public" / "artefatos"
INDEX = PROJECT / "docs" / "source-index.json"

# Positions reconstructed from the local question TXT files; the matching lesson
# PDFs are used whenever they explicitly establish the answer. They are ordered
# by original TXT filename.
QUIZ_ANSWER_POSITIONS = {
    "Back-End": [4, 1, 4, 2, 2, 3, 3, 4, 3, 2, 4, 1, 4, 4, 2, 2, 1, 4, 4, 4, 1],
    "Banco de dados": [1, 1, 4, 3, 3, 3, 3, 4, 3, 1, 4, 1, 2, 1, 4, 2, 3, 3, 1, 4, 4],
    "Front-End": [2, 2, 4, 2, 1, 3, 1, 1, 4, 2, 3, 2, 4, 2, 4, 3, 3, 3, 3, 3, 3],
    "Inteligência Artificial": [1, 4, 3, 1, 2, 3, 2, 4, 3, 3, 2, 4, 4, 4, 1, 4, 1, 3],
    "Programação Mobile": [4, 3, 3, 2, 3, 3, 4, 3, 4, 2, 3, 2, 3, 2, 3, 3, 3, 3, 2, 3, 1],
    "Projeto Multidisciplinar": [3, 3, 3, 4, 2, 4, 3, 3, 3, 3, 4, 1, 4, 2, 2, 3, 4, 3, 1, 3, 4],
    "Versionamento de Código": [2, 2, 2, 1, 1, 3, 4, 1, 1, 2, 3, 1, 4, 3, 3, 1, 4, 1],
}

# Answers are present only when the roteiro supplies enough fixed information.
# Execution- and data-dependent work stays explicitly practical.
WRITTEN_ANSWERS = {
    "47278": "Autenticação confirma a identidade de quem chama a API, por exemplo com JWT. Autorização verifica o que essa identidade pode acessar. HTTPS protege os dados em trânsito contra interceptação, e a criptografia em repouso reduz a exposição se o armazenamento for comprometido.",
    "47280": "Teste de penetração simula tentativas reais de ataque antes do lançamento. Ferramentas como OWASP ZAP ajudam a localizar falhas em endpoints; a verificação deve cobrir injeção de SQL, XSS e exposição de dados sensíveis. Testes contínuos detectam regressões de segurança a cada mudança.",
    "47217": "Antes de DELETE ou UPDATE, confirme o alvo com uma consulta SELECT usando a mesma condição WHERE, faça backup ou trabalhe em transação quando possível e evite comandos sem WHERE. Assim a alteração pode ser conferida e, se necessário, revertida.",
    "83039": "Moderador em tempo real: arquitetura reativa, pois precisa responder em poucos milissegundos. Diagnóstico médico: deliberativa, porque exige análise de hipóteses e justificativas. Robô de armazém: híbrida, unindo reação imediata a obstáculos e planejamento de rota. A escolha depende de latência, explicabilidade, precisão e planejamento exigidos.",
    "83062": "Correção de arredondamento altera PATCH: 1.2.3 para 1.2.4. A nova função compatível altera MINOR: 1.3.0. Exigir o novo parâmetro quebra clientes antigos, portanto altera MAJOR: 2.0.0. Tags marcam referências imutáveis de release; mudar MAJOR sem necessidade dificulta a leitura de compatibilidade.",
    "83067": "Changelog v1.3.0: Features: cálculo de imposto e exportação em PDF. Correções: arredondamento. Melhorias internas: refatoração de organização. Documentação: atualização do README. É MINOR porque acrescenta funcionalidades compatíveis, sem quebra de API. Mensagens padronizadas permitem classificar e gerar esse documento automaticamente.",
    "82907": "ACK manual só deve ocorrer depois do processamento concluído, pois confirmar antes pode perder a mensagem em caso de falha. NACK não precisa apagar a mensagem: ela pode ser reenfileirada ou isolada. Falhas temporárias justificam nova tentativa; falhas definitivas devem seguir para uma dead-letter queue. Use direct para uma fila específica, fanout para todos os serviços inscritos e topic para padrões como pedido.*.",
    "82912": "Quando a entrada supera a saída e há muitas mensagens não confirmadas, o indício principal é consumo lento ou sobrecarregado. Um pré-fetch alto pode concentrar mensagens em um consumidor; reduzi-lo ajuda a distribuir a carga, mas não substitui investigar o processamento e a quantidade de consumidores. A Management UI permite relacionar filas, taxas, consumidores e confirmações; direct atende uma fila, fanout replica para várias e topic usa padrões de chave.",
}


def normalize(value: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.replace("\r", "").split("\n")]
    compact: list[str] = []
    for line in lines:
        if line or (compact and compact[-1]):
            compact.append(line)
    return "\n".join(compact).strip()


def slugify(value: str) -> str:
    value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii").lower()
    return re.sub(r"-+", "-", re.sub(r"[^a-z0-9]+", "-", value)).strip("-")


def number_from(value: str, label: str) -> int:
    match = re.search(rf"{label}\s*(\d+)", value, re.IGNORECASE)
    if not match:
        raise RuntimeError(f"Não foi possível localizar {label} em: {value}")
    return int(match.group(1))


def read_docx(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return normalize("\n\n".join(blocks))


def title_from(text: str, fallback: str) -> str:
    for pattern in (r"(?:Título da atividade|Atividade)\s*:\s*([^\n]+)", r"Título\s*:\s*([^\n]+)"):
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return normalize(match.group(1))
    return re.sub(r"\s*-\s*\d+$", "", fallback.replace("Atividade Pratica - ", "")).strip()


def source_path(path: Path) -> str:
    return path.relative_to(PROJECT.parent).as_posix()


def write_mdx(data: dict, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(f"---\n{json.dumps(data, ensure_ascii=False)}\n---\n", encoding="utf-8")


def reset_generated_directory(path: Path) -> None:
    resolved = path.resolve()
    allowed = {CONTENT.resolve(), PUBLIC_ARTIFACTS.resolve()}
    if resolved not in allowed or PROJECT.resolve() not in resolved.parents:
        raise RuntimeError(f"Diretório gerado inválido: {path}")
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def quiz_from_txt(path: Path, answer_position: int) -> dict:
    lines = [line.strip() for line in path.read_text(encoding="utf-8-sig").splitlines()]
    try:
        start = next(index for index, line in enumerate(lines) if line.startswith("Fonte:")) + 1
    except StopIteration as error:
        raise RuntimeError(f"Fonte ausente em {path}") from error
    payload = [line for line in lines[start:] if line]
    if len(payload) < 3:
        raise RuntimeError(f"Questão incompleta em {path}")
    question, alternatives = payload[0], payload[1:]
    if not 1 <= answer_position <= len(alternatives):
        raise RuntimeError(f"Posição de resposta inválida em {path}")
    return {
        "pergunta": question,
        "alternativas": alternatives,
        "resposta": alternatives[answer_position - 1],
        "fonte": source_path(path),
    }


def main() -> None:
    if not SOURCE.is_dir():
        raise RuntimeError(f"Fonte não encontrada: {SOURCE}")
    reset_generated_directory(CONTENT)
    reset_generated_directory(PUBLIC_ARTIFACTS)
    records: list[dict] = []
    activity_count = quiz_page_count = question_count = practical_count = written_count = 0

    for subject_dir in sorted(path for path in SOURCE.iterdir() if path.is_dir()):
        subject = subject_dir.name
        subject_slug = slugify(subject)
        term_dirs = [path for path in subject_dir.iterdir() if path.is_dir() and path.name.startswith("3")]
        if len(term_dirs) != 1:
            raise RuntimeError(f"Bimestre 3 ambíguo para {subject}: {term_dirs}")
        quiz_files_by_week: dict[int, list[Path]] = defaultdict(list)
        week_dirs = sorted((path for path in term_dirs[0].iterdir() if path.is_dir()), key=lambda path: number_from(path.name, "Semana"))
        for week_dir in week_dirs:
            week = number_from(week_dir.name, "Semana")
            activity_files = sorted(week_dir.glob("Atividade Pratica*.docx"), key=lambda path: number_from(path.name, "Aula"))
            for activity_file in activity_files:
                order = number_from(activity_file.name, "Aula")
                text = read_docx(activity_file)
                source_id = re.search(r"(\d+)\.docx$", activity_file.name)
                fixed_answer = WRITTEN_ANSWERS.get(source_id.group(1) if source_id else "")
                activity = {
                    "tipo": "resposta" if fixed_answer else "pratica",
                    "enunciado": text,
                    "resposta": fixed_answer or "O roteiro não traz uma resposta escrita pronta. A entrega solicitada é prática: siga o procedimento original e registre a sua execução.",
                    "fonte": source_path(activity_file),
                }
                data = {"materia": subject, "materiaSlug": subject_slug, "bimestre": 3, "semana": week, "ordem": order, "aula": f"a{order}", "titulo": title_from(text, activity_file.stem), "atividade": activity, "quizzes": []}
                output = CONTENT / subject_slug / f"semana-{week}" / f"a{order}.mdx"
                write_mdx(data, output)
                records.append({"tipo": "atividade", "fonte": activity["fonte"], "pagina": output.relative_to(PROJECT).as_posix(), "rota": f"/{subject_slug}/3/semana-{week}/a{order}"})
                activity_count += 1
                written_count += bool(fixed_answer)
                practical_count += not bool(fixed_answer)
            quiz_files_by_week[week].extend(sorted(week_dir.glob("Pause e Responda*.txt")))

        positions = QUIZ_ANSWER_POSITIONS.get(subject)
        quiz_files = [file for week in sorted(quiz_files_by_week) for file in quiz_files_by_week[week]]
        if positions is None or len(positions) != len(quiz_files):
            raise RuntimeError(f"Mapa de respostas incompleto para {subject}: {len(positions or [])}/{len(quiz_files)}")
        grouped_quizzes: dict[int, list[dict]] = defaultdict(list)
        for quiz_file, answer_position in zip(quiz_files, positions, strict=True):
            week = number_from(quiz_file.parent.name, "Semana")
            grouped_quizzes[week].append(quiz_from_txt(quiz_file, answer_position))
            question_count += 1
        for week, quizzes in grouped_quizzes.items():
            activities_in_week = list((CONTENT / subject_slug / f"semana-{week}").glob("a*.mdx"))
            data = {"materia": subject, "materiaSlug": subject_slug, "bimestre": 3, "semana": week, "ordem": len(activities_in_week) + 1, "aula": "q", "titulo": "Pause e Responda", "quizzes": quizzes}
            output = CONTENT / subject_slug / f"semana-{week}" / "q.mdx"
            write_mdx(data, output)
            records.extend({"tipo": "pause", "fonte": quiz["fonte"], "pagina": output.relative_to(PROJECT).as_posix(), "rota": f"/{subject_slug}/3/semana-{week}/q"} for quiz in quizzes)
            quiz_page_count += 1

    INDEX.parent.mkdir(parents=True, exist_ok=True)
    INDEX.write_text(json.dumps({"fonte": "Material", "registros": records}, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Atividades: {activity_count}; práticas: {practical_count}; respostas escritas: {written_count}; pausas: {quiz_page_count}; questões: {question_count}")


if __name__ == "__main__":
    main()
